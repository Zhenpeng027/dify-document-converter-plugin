"""
文档转换核心模块
支持多种格式的文档转换
"""

import os
import re
from typing import List, Dict, Any, Optional
from .document_styles import DocumentStyleManager, FontStyle, ParagraphStyle
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
import markdown
from bs4 import BeautifulSoup

class DocumentConverter:
    """文档转换器"""
    
    def __init__(self, style_manager: DocumentStyleManager):
        self.style_manager = style_manager
    
    def convert_markdown_to_docx(self, markdown_content: str, output_path: str = None) -> Document:
        """将Markdown转换为Word文档"""
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置页面样式
            self._apply_page_style(doc)
            
            # 解析Markdown内容
            lines = markdown_content.split('\n')
            current_paragraph = []
            in_code_block = False
            code_language = ""
            
            for line in lines:
                # 处理代码块
                if line.strip().startswith('```'):
                    if in_code_block:
                        # 结束代码块
                        if current_paragraph:
                            code_content = '\n'.join(current_paragraph)
                            self._add_code_block(doc, code_content, code_language)
                            current_paragraph = []
                        in_code_block = False
                        code_language = ""
                    else:
                        # 开始代码块
                        in_code_block = True
                        code_language = line.strip()[3:].strip()
                    continue
                
                if in_code_block:
                    current_paragraph.append(line)
                    continue
                
                # 处理标题
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()
                    if level == 1:
                        self._add_heading(doc, text, 'title')
                    elif level <= 3:
                        self._add_heading(doc, text, f'heading{level}')
                    else:
                        self._add_paragraph(doc, text, 'normal')
                elif line.strip():
                    # 普通段落
                    self._add_paragraph(doc, line, 'normal')
                else:
                    # 空行 - 添加间距
                    doc.add_paragraph()
            
            # 处理剩余的代码块
            if in_code_block and current_paragraph:
                code_content = '\n'.join(current_paragraph)
                self._add_code_block(doc, code_content, code_language)
            
            # 添加页眉页脚
            self._add_header_footer(doc)
            
            # 保存文档
            if output_path:
                doc.save(output_path)
                return f"文档已保存到: {output_path}"
            else:
                return doc
                
        except Exception as e:
            return f"转换失败: {str(e)}"
    
    def convert_text_to_docx(self, text_content: str, output_path: str) -> str:
        """将纯文本转换为Word文档"""
        try:
            doc = Document()
            self._apply_page_style(doc)
            
            # 按行处理文本
            lines = text_content.split('\n')
            for line in lines:
                if line.strip():
                    self._add_paragraph(doc, line, 'normal')
                else:
                    doc.add_paragraph()
            
            self._add_header_footer(doc)
            doc.save(output_path)
            return f"文档已保存到: {output_path}"
            
        except Exception as e:
            return f"转换失败: {str(e)}"
    
    def _apply_page_style(self, doc: Document):
        """应用页面样式"""
        section = doc.sections[0]
        page_style = self.style_manager.page_style
        
        # 设置页面尺寸和边距
        section.page_width = Inches(page_style.width / 25.4)  # mm to inches
        section.page_height = Inches(page_style.height / 25.4)
        section.top_margin = Inches(page_style.margin_top / 25.4)
        section.bottom_margin = Inches(page_style.margin_bottom / 25.4)
        section.left_margin = Inches(page_style.margin_left / 25.4)
        section.right_margin = Inches(page_style.margin_right / 25.4)
        
        # 设置页面方向
        if page_style.orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
    
    def _add_heading(self, doc: Document, text: str, style_name: str):
        """添加标题"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        
        style_config = self.style_manager.get_style(style_name)
        if style_config:
            self._apply_font_style(run, style_config['font'])
            self._apply_paragraph_style(paragraph, style_config['paragraph'])
    
    def _add_paragraph(self, doc: Document, text: str, style_name: str):
        """添加段落"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        
        style_config = self.style_manager.get_style(style_name)
        if style_config:
            self._apply_font_style(run, style_config['font'])
            self._apply_paragraph_style(paragraph, style_config['paragraph'])
    
    def _add_code_block(self, doc: Document, code: str, language: str = ""):
        """添加代码块"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(code)
        
        style_config = self.style_manager.get_style('code')
        if style_config:
            self._apply_font_style(run, style_config['font'])
            self._apply_paragraph_style(paragraph, style_config['paragraph'])
    
    def _apply_font_style(self, run, font_style: FontStyle):
        """应用字体样式"""
        if font_style:
            run.font.name = font_style.family
            run.font.size = Pt(font_style.size)
            run.font.bold = font_style.bold
            run.font.italic = font_style.italic
            
            # 解析颜色
            if font_style.color.startswith('#'):
                color_hex = font_style.color[1:]
                try:
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except ValueError:
                    pass  # 忽略无效颜色
    
    def _apply_paragraph_style(self, paragraph, paragraph_style: ParagraphStyle):
        """应用段落样式"""
        if paragraph_style:
            # 对齐方式
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            paragraph.alignment = alignment_map.get(paragraph_style.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            
            # 行间距
            paragraph.paragraph_format.line_spacing = paragraph_style.line_spacing
            
            # 段前段后间距
            paragraph.paragraph_format.space_before = Pt(paragraph_style.space_before)
            paragraph.paragraph_format.space_after = Pt(paragraph_style.space_after)
            
            # 缩进
            if paragraph_style.indent_first_line:
                paragraph.paragraph_format.first_line_indent = Pt(paragraph_style.indent_first_line)
            if paragraph_style.indent_left:
                paragraph.paragraph_format.left_indent = Pt(paragraph_style.indent_left)
            if paragraph_style.indent_right:
                paragraph.paragraph_format.right_indent = Pt(paragraph_style.indent_right)
    
    def _add_header_footer(self, doc: Document):
        """添加页眉页脚"""
        section = doc.sections[0]
        
        # 添加页眉
        if self.style_manager.header_style.enabled and self.style_manager.header_style.content:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = self.style_manager.header_style.content
            
            # 应用页眉样式
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT
            }
            header_para.alignment = alignment_map.get(
                self.style_manager.header_style.alignment, 
                WD_ALIGN_PARAGRAPH.CENTER
            )
            
            # 应用字体样式
            if self.style_manager.header_style.font:
                run = header_para.runs[0] if header_para.runs else header_para.add_run()
                self._apply_font_style(run, self.style_manager.header_style.font)
        
        # 添加页脚
        if self.style_manager.footer_style.enabled and self.style_manager.footer_style.content:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = self.style_manager.footer_style.content
            
            # 应用页脚样式
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT
            }
            footer_para.alignment = alignment_map.get(
                self.style_manager.footer_style.alignment, 
                WD_ALIGN_PARAGRAPH.CENTER
            )
            
            # 应用字体样式
            if self.style_manager.footer_style.font:
                run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
                self._apply_font_style(run, self.style_manager.footer_style.font)