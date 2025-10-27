"""
Markdown to Word Converter Tool for Dify
"""

import os
import tempfile
import json
from typing import Any, Dict, List, Union
from dify_plugin.entities.tool import ToolInvokeMessage
from dify_plugin.tools.tool import Tool
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
import markdown
from bs4 import BeautifulSoup


class MarkdownToWordTool(Tool):
    """
    Markdown to Word converter tool
    """
    
    def _invoke(self, user_id: str, tool_parameters: Dict[str, Any]) -> Union[ToolInvokeMessage, List[ToolInvokeMessage]]:
        """
        Invoke the markdown to word conversion tool
        
        Args:
            user_id: The user ID
            tool_parameters: Tool parameters including markdown content and style options
            
        Returns:
            ToolInvokeMessage with the converted Word document
        """
        try:
            # Extract parameters
            markdown_content = tool_parameters.get('markdown_content', '')
            style_template = tool_parameters.get('style_template', '学术论文')
            custom_styles_json = tool_parameters.get('custom_styles', '')
            
            if not markdown_content:
                return ToolInvokeMessage(type='text', message='Error: Markdown content is required')
            
            # Parse custom styles JSON if provided
            custom_styles = {}
            if custom_styles_json:
                try:
                    custom_styles = json.loads(custom_styles_json)
                except json.JSONDecodeError as e:
                    return ToolInvokeMessage(type='text', message=f'Error: Invalid JSON in custom_styles: {str(e)}')
            
            # Initialize style manager with template
            style_manager = self._create_style_manager(style_template, custom_styles)
            
            # Convert markdown to Word
            doc = self._convert_markdown_to_docx(markdown_content, style_manager)
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            # Read file content
            with open(temp_file.name, 'rb') as f:
                file_content = f.read()
            
            # Clean up
            os.unlink(temp_file.name)
            
            return ToolInvokeMessage(
                type='blob',
                message=file_content,
                meta={
                    'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'filename': 'converted_document.docx'
                }
            )
            
        except Exception as e:
            return ToolInvokeMessage(type='text', message=f'Error converting document: {str(e)}')
    
    def _create_style_manager(self, template_name: str, custom_styles: Dict[str, Any]):
        """Create and configure style manager"""
        from ..document_styles import DocumentStyleManager, FontStyle, ParagraphStyle, PageStyle, STYLE_TEMPLATES
        
        style_manager = DocumentStyleManager()
        
        # Apply template if exists
        if template_name in STYLE_TEMPLATES:
            template = STYLE_TEMPLATES[template_name]
            for style_name, style_config in template.items():
                if style_name == 'page':
                    style_manager.set_page_style(style_config)
                else:
                    font_style = style_config.get('font')
                    paragraph_style = style_config.get('paragraph')
                    style_manager.update_style(style_name, font_style, paragraph_style)
        
        # Apply custom styles
        for style_name, style_config in custom_styles.items():
            font_config = style_config.get('font', {})
            paragraph_config = style_config.get('paragraph', {})
            
            font_style = FontStyle(**font_config) if font_config else None
            paragraph_style = ParagraphStyle(**paragraph_config) if paragraph_config else None
            
            style_manager.update_style(style_name, font_style, paragraph_style)
        
        return style_manager
    
    def _convert_markdown_to_docx(self, markdown_content: str, style_manager) -> Document:
        """Convert markdown content to Word document"""
        from ..document_converter import DocumentConverter
        
        converter = DocumentConverter(style_manager)
        
        # Create Word document
        doc = Document()
        
        # Apply page style
        self._apply_page_style(doc, style_manager)
        
        # Parse and convert markdown content
        lines = markdown_content.split('\n')
        current_paragraph = []
        in_code_block = False
        code_language = ""
        
        for line in lines:
            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End code block
                    if current_paragraph:
                        code_content = '\n'.join(current_paragraph)
                        self._add_code_block(doc, code_content, code_language, style_manager)
                        current_paragraph = []
                    in_code_block = False
                    code_language = ""
                else:
                    # Start code block
                    in_code_block = True
                    code_language = line.strip()[3:].strip()
                continue
            
            if in_code_block:
                current_paragraph.append(line)
                continue
            
            # Handle headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if level == 1:
                    self._add_heading(doc, text, 'title', style_manager)
                elif level <= 3:
                    self._add_heading(doc, text, f'heading{level}', style_manager)
                else:
                    self._add_paragraph(doc, text, 'normal', style_manager)
            elif line.strip():
                # Regular paragraph
                self._add_paragraph(doc, line, 'normal', style_manager)
            else:
                # Empty line - add spacing
                doc.add_paragraph()
        
        # Handle remaining code block
        if in_code_block and current_paragraph:
            code_content = '\n'.join(current_paragraph)
            self._add_code_block(doc, code_content, code_language, style_manager)
        
        return doc
    
    def _apply_page_style(self, doc: Document, style_manager):
        """Apply page style to document"""
        section = doc.sections[0]
        page_style = style_manager.page_style
        
        section.page_width = Inches(page_style.width / 25.4)
        section.page_height = Inches(page_style.height / 25.4)
        section.top_margin = Inches(page_style.margin_top / 25.4)
        section.bottom_margin = Inches(page_style.margin_bottom / 25.4)
        section.left_margin = Inches(page_style.margin_left / 25.4)
        section.right_margin = Inches(page_style.margin_right / 25.4)
        
        if page_style.orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
    
    def _add_heading(self, doc: Document, text: str, style_name: str, style_manager):
        """Add heading with style"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        
        style_config = style_manager.get_style(style_name)
        if style_config:
            self._apply_font_style(run, style_config['font'])
            self._apply_paragraph_style(paragraph, style_config['paragraph'])
    
    def _add_paragraph(self, doc: Document, text: str, style_name: str, style_manager):
        """Add paragraph with style"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        
        style_config = style_manager.get_style(style_name)
        if style_config:
            self._apply_font_style(run, style_config['font'])
            self._apply_paragraph_style(paragraph, style_config['paragraph'])
    
    def _add_code_block(self, doc: Document, code: str, language: str, style_manager):
        """Add code block with style"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(code)
        
        style_config = style_manager.get_style('code')
        if style_config:
            self._apply_font_style(run, style_config['font'])
            self._apply_paragraph_style(paragraph, style_config['paragraph'])
    
    def _apply_font_style(self, run, font_style):
        """Apply font style to run"""
        if font_style:
            run.font.name = font_style.family
            run.font.size = Pt(font_style.size)
            run.font.bold = font_style.bold
            run.font.italic = font_style.italic
            
            # Parse color
            if font_style.color.startswith('#'):
                color_hex = font_style.color[1:]
                r = int(color_hex[0:2], 16)
                g = int(color_hex[2:4], 16)
                b = int(color_hex[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
    
    def _apply_paragraph_style(self, paragraph, paragraph_style):
        """Apply paragraph style"""
        if paragraph_style:
            # Alignment
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            paragraph.alignment = alignment_map.get(paragraph_style.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            
            # Line spacing
            paragraph.paragraph_format.line_spacing = paragraph_style.line_spacing
            
            # Spacing
            paragraph.paragraph_format.space_before = Pt(paragraph_style.space_before)
            paragraph.paragraph_format.space_after = Pt(paragraph_style.space_after)
            
            # Indentation
            if paragraph_style.indent_first_line:
                paragraph.paragraph_format.first_line_indent = Pt(paragraph_style.indent_first_line)
            if paragraph_style.indent_left:
                paragraph.paragraph_format.left_indent = Pt(paragraph_style.indent_left)
            if paragraph_style.indent_right:
                paragraph.paragraph_format.right_indent = Pt(paragraph_style.indent_right)